# scripts/generate_docs.py
"""
Gera documentos corporativos sintéticos (.docx e .txt) para popular o RAG.
Sem dependências pesadas: usa apenas python-docx (já está no seu projeto).

Uso:
  python scripts/generate_docs.py --out data/docs --n 12
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import random
import argparse
import textwrap
import datetime as dt

CATEGORIES = {
    "Políticas de Segurança": [
        "Política de Senhas",
        "Diretrizes de Acesso Remoto (VPN)",
        "Classificação e Manuseio de Informações",
        "Backup e Recuperação de Desastres",
        "Uso de Antivírus e Atualizações"
    ],
    "TI & Suporte": [
        "Abertura de Chamados no Service Desk",
        "Procedimentos de Atendimento Técnico",
        "Política de Uso de E-mail Corporativo",
        "Padrão de Nomenclatura de Arquivos",
        "Checklist de Entrega/Devolução de Equipamentos"
    ],
    "RH & Onboarding": [
        "Guia de Integração do Colaborador",
        "Benefícios e Políticas Internas",
        "Boas Práticas de Home Office",
        "Conduta e Ética",
        "Comunicação e Canais Oficiais"
    ],
    "Compliance & LGPD": [
        "Política de Privacidade e Proteção de Dados",
        "Direitos do Titular de Dados",
        "Incidentes de Segurança e Notificação",
        "Retenção e Eliminação de Dados",
        "Termos de Uso de Sistemas"
    ],
}

BULLETS = [
    "Cumprir as diretrizes descritas neste documento.",
    "Em caso de dúvida, contatar o responsável pelo processo.",
    "Manter registros atualizados e acessíveis para auditoria.",
    "Respeitar prazos e fluxos de aprovação.",
    "Reportar incidentes ou desvios imediatamente."
]

FAQ_QA = [
    ("Como solicitar acesso a um sistema?",
     "Abra um chamado no Service Desk anexando a aprovação do gestor."),
    ("Quais são os requisitos de senha?",
     "Mínimo de 12 caracteres, com letras maiúsculas/minúsculas, números e símbolo."),
    ("Qual o prazo para resposta de chamados críticos?",
     "Até 4 horas úteis, com prioridade 1 na fila de atendimento."),
    ("Posso usar e-mail pessoal para trabalho?",
     "Não. Utilize apenas o e-mail corporativo para assuntos da empresa."),
]

PARAGRAPH = (
    "Este documento estabelece normas e orientações para garantir o bom "
    "funcionamento e a segurança dos processos internos. Seu objetivo é "
    "padronizar procedimentos, mitigar riscos e assegurar conformidade com "
    "regulamentos aplicáveis (ex.: LGPD)."
)

def hdr(doc, title, subtitle):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.runs[0].font.size = Pt(11)
        p2.runs[0].italic = True

def section(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

def para(doc, text):
    for chunk in textwrap.wrap(text, width=100):
        doc.add_paragraph(chunk)

def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def mini_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.cell(i, j).text = cell

def faq(doc, qa):
    section(doc, "FAQ")
    for q, a in qa:
        doc.add_paragraph(f"Q: {q}", style=None).runs[0].bold = True
        doc.add_paragraph(f"A: {a}")

def make_docx(out_dir: Path, category: str, title: str):
    doc = Document()
    hdr(doc, f"{title}", f"Categoria: {category} – Versão {random.randint(1,5)}.{random.randint(0,9)}")
    doc.add_paragraph(f"Data: {dt.date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph()

    section(doc, "1. Objetivo")
    para(doc, PARAGRAPH)

    section(doc, "2. Escopo")
    para(doc, "Aplica-se a todos os colaboradores, estagiários e terceiros com acesso aos recursos da organização.")

    section(doc, "3. Diretrizes")
    bullets(doc, random.sample(BULLETS, k=3))
    doc.add_paragraph()
    mini_table(doc, [
        ["Responsável", "E-mail", "SLA"],
        ["Service Desk", "suporte@empresa.com", "8x5"],
        ["Segurança da Informação", "si@empresa.com", "24x7 (incidentes)"],
    ])

    faq(doc, random.sample(FAQ_QA, k=min(3, len(FAQ_QA))))
    doc.add_paragraph()
    section(doc, "4. Conformidade")
    para(doc, "O não cumprimento destas diretrizes pode acarretar medidas disciplinares e registro de não conformidade.")

    out_file = out_dir / f"{sanitize(title)}.docx"
    doc.save(out_file)
    return out_file

def make_txt(out_dir: Path, category: str, title: str):
    content = [
        f"# {title}",
        f"_Categoria: {category}_",
        "",
        "## Objetivo",
        PARAGRAPH,
        "",
        "## Diretrizes",
        "- " + "\n- ".join(random.sample(BULLETS, k=3)),
        "",
        "## FAQ",
    ]
    for q, a in random.sample(FAQ_QA, k=min(3, len(FAQ_QA))):
        content.append(f"Q: {q}")
        content.append(f"A: {a}")
        content.append("")
    out_file = out_dir / f"{sanitize(title)}.txt"
    out_file.write_text("\n".join(content), encoding="utf-8")
    return out_file

def sanitize(name: str) -> str:
    return "".join(c for c in name if c.isalnum() or c in " _-").strip().replace(" ", "_")

def generate(n: int, out: Path):
    out.mkdir(parents=True, exist_ok=True)
    all_titles = []
    for cat, titles in CATEGORIES.items():
        all_titles += [(cat, t) for t in titles]
    if n > len(all_titles):
        # repete com sufixo
        extra = n - len(all_titles)
        for i in range(extra):
            cat, t = random.choice(list(CATEGORIES.items()))
        # nada crítico — repetição aleatória
    selected = random.sample(all_titles, k=min(n, len(all_titles))) if n <= len(all_titles) else all_titles

    created = []
    for cat, title in selected:
        created.append(make_docx(out, cat, title))
        created.append(make_txt(out, cat, title))
    return created

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default="data/docs", help="Pasta de saída (default: data/docs)")
    ap.add_argument("--n", type=int, default=10, help="Quantidade de documentos (títulos base)")
    args = ap.parse_args()

    out = Path(args.out)
    files = generate(args.n, out)
    print(f"OK! Gerados {len(files)} arquivos em {out.resolve()}")
    for f in files:
        print(" -", f)

if __name__ == "__main__":
    main()
